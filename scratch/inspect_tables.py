import docx

def inspect_tables(file_path):
    doc = docx.Document(file_path)
    print(f"Total tables: {len(doc.tables)}")
    for i, table in enumerate(doc.tables[:5]): # Inspect first 5 tables
        print(f"\nTable {i}:")
        for row in table.rows[:3]: # First 3 rows
            cells = [cell.text.strip() for cell in row.cells]
            print(f"  {cells}")

if __name__ == "__main__":
    inspect_tables("Beginner vocabulary last one.docx")
